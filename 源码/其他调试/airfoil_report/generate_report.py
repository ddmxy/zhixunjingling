#!/usr/bin/env python3
"""Generate experiment report sections 4 (fill-in) and 5 with expanded analysis."""
import os

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

DESKTOP = os.path.join(os.environ["USERPROFILE"], "Desktop")
CSV = os.path.join(DESKTOP, "翼型实验_系数表.csv")

# 实验当天环境参数（风洞实验室常规取值；若实验时实测了温湿度计/气压计，以实测为准）
P_ATM = 101325  # Pa，标准大气压
T_C = 25.0  # °C，实验室室温
R_AIR = 287.0  # J/(kg·K)
MU_25C = 1.84e-5  # Pa·s，25°C 空气动力粘度
V_INF = 20.0  # m/s
CHORD = 0.2  # m

T_K = T_C + 273.15
RHO = P_ATM / (R_AIR * T_K)
Q_INF = 0.5 * RHO * V_INF**2
RE = RHO * V_INF * CHORD / MU_25C


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, indent=True):
    p = doc.add_paragraph(text)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = "宋体"


def add_image(doc, filename, caption, width_cm=14):
    path = os.path.join(DESKTOP, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width_cm))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(10.5)
            run.font.name = "宋体"


def main():
    df = pd.read_csv(CSV)
    doc = Document()

    # ========== 第4节 实验步骤（填空） ==========
    add_heading(doc, "4．实验步骤介绍（实验条件填写）", level=1)

    add_para(
        doc,
        "本次实验在低速低湍流风洞中进行，实验条件按指导书要求记录如下。"
        "其中大气压和气温为风洞实验室环境参数，用于计算空气密度和雷诺数；"
        "若实验现场有温湿度计、气压计读数，应以实测值替换下表中的参考值。"
    )

    cond_table = doc.add_table(rows=5, cols=2)
    cond_table.style = "Table Grid"
    rows = [
        ("风速 V / (m·s⁻¹)", f"{V_INF:.0f}"),
        ("攻角 α / (°)", "−2，0，4，8，12，16"),
        ("大气压 P / Pa", f"{P_ATM:.0f}"),
        ("气温 t / °C", f"{T_C:.1f}"),
        ("备注", "NACA0025 翼型，弦长 c = 0.2 m"),
    ]
    for i, (k, v) in enumerate(rows):
        cond_table.rows[i].cells[0].text = k
        cond_table.rows[i].cells[1].text = v

    cap4 = doc.add_paragraph("表0  实验条件记录表")
    cap4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "4.1 大气压与气温如何填写", level=2)
    add_para(
        doc,
        f"（1）大气压 P：实验在室内风洞实验室进行，若无专门气压记录，"
        f"可取标准大气压 P = {P_ATM:.0f} Pa（即 101.325 kPa）。"
        f"若实验室墙上有气压计或天气预报当日气压约为 1010～1020 hPa，"
        f"也可换算为 Pa 后填入（1 hPa = 100 Pa）。"
    )
    add_para(
        doc,
        f"（2）气温 t：取实验进行时的实验室室温。"
        f"本实验参考取 t = {T_C:.0f} °C（夏季/春秋室内空调环境常见值）。"
        f"若实验时读取了温度计，应填写实测值，例如 t = 23 °C。"
    )
    add_para(
        doc,
        f"（3）由状态方程计算空气密度：ρ = P / (R·T)，"
        f"取 R = 287 J/(kg·K)，T = {T_K:.2f} K，得 ρ = {RHO:.3f} kg/m³。"
        f"来流动压 q∞ = ½ρV² = {Q_INF:.1f} Pa。"
        f"雷诺数 Re = ρ·V·c / μ ≈ {RE:.2e}（取 μ ≈ {MU_25C:.2e} Pa·s）。"
    )

    add_heading(doc, "4.2 实验步骤简述", level=2)
    steps = [
        "打开计算机桌面数据采集系统，进入翼型测压实验界面；",
        f"在测试界面输入风速 V = {V_INF:.0f} m/s 及当前攻角 α；",
        "设置齿轮位置与测压孔编号对应关系，多通道扫描微压阀与测压孔一一连接；",
        "系统自动采集各测压孔静压，经换算得到上下表面压强系数 Cp 并保存；",
        "依次改变攻角为 −2°、0°、4°、8°、12°、16°，重复采集，得到各工况 Excel 数据文件。",
    ]
    for i, s in enumerate(steps, 1):
        add_para(doc, f"步骤{i}：{s}")

    doc.add_page_break()

    # ========== 第5节 实验结果及处理 ==========
    add_heading(doc, "5．实验结果及处理", level=1)

    add_heading(doc, "5.1 实验条件与数据处理说明", level=2)
    add_para(
        doc,
        f"翼型为 NACA0025 对称翼型，弦长 c = {CHORD} m，相对厚度 25%，"
        f"表面沿弦向布置 25 个测压孔。来流速度 V∞ = {V_INF} m/s，"
        f"大气压 P = {P_ATM} Pa，气温 t = {T_C} °C，"
        f"空气密度 ρ = {RHO:.3f} kg/m³，动压 q∞ = {Q_INF:.1f} Pa，"
        f"雷诺数 Re ≈ {RE:.2e}。"
        f"攻角取 −2°、0°、4°、8°、12°、16° 共六组工况。"
    )
    add_para(
        doc,
        "压强系数 Cp = (p − p∞) / q∞，其中 p 为测压孔处静压，p∞ 为风洞来流静压。"
        "由上下表面 Cp 沿弦向分布，经数值积分得到法向力系数 Cn、轴向力系数 Ca，"
        "再按坐标变换得到升力系数 Cl 和阻力系数 Cd。"
    )

    add_heading(doc, "5.2 翼型表面压强分布", level=2)
    add_para(
        doc,
        "图1为 α = 8° 时 NACA0025 翼型上下表面压强系数沿弦向 x/c 的分布。"
        "可以看出：在前缘附近（x/c < 0.1），上表面 Cp 迅速降低，出现明显吸力峰"
        "（最低 Cp 约 −1.9），表明气流加速、静压显著低于来流；"
        "下表面 Cp 相对较高甚至为正，上下压差最大，是升力的主要来源。"
        "沿弦向后，上表面负压逐渐减小，下表面压强也趋于恢复，"
        "在尾缘附近上下表面 Cp 趋于汇合（Kutta 条件）。"
    )
    add_para(
        doc,
        "图2对比了各攻角下的 Cp 分布。随 α 增大，上表面负压区向前缘扩展、吸力增强，"
        "下表面正压区扩大，上下积分面积增大，对应升力系数升高。"
        "当 α = 12° 时上表面负压峰值仍较大；至 α = 16° 时，"
        "上表面前缘吸力峰明显减弱、负压区缩小，预示流动分离加剧、接近失速。"
    )
    add_image(doc, "图1_压强分布_alpha8.png", "图1  NACA0025 翼型压强分布（α = 8°，V = 20 m/s）")
    add_image(doc, "图2_各攻角压强分布.png", "图2  各攻角下翼型表面压强系数分布", width_cm=15)

    add_heading(doc, "5.3 升力系数与阻力系数的计算", level=2)
    add_para(doc, "（1）法向力系数（垂直于弦向）：")
    add_para(doc, "Cn = ∫₀¹ (Cpl − Cpu) d(x/c)", indent=False)
    add_para(
        doc,
        "式中 Cpl、Cpu 分别为下表面、上表面压强系数。"
        "采用梯形法则对 25 个测压点对应的 x/c 进行数值积分。"
        "Cn 在物理意义上等于上下表面 Cp 差曲线下的面积，"
        "直接反映翼型所受法向力（单位展长、无量纲化）。"
    )
    add_para(doc, "（2）轴向力系数（沿弦向，压差阻力相关）：")
    add_para(doc, "Ca = ∫ (Cpl·dyl/dx − Cpu·dyu/dx) d(x/c)", indent=False)
    add_para(
        doc,
        "其中 dyu/dx、dyl/dx 由 NACA0025 厚度分布公式对 x/c 求导得到。"
        "前缘 x/c < 0.025 区域因几何导数奇异，不参与 Ca 积分（实验常规处理）。"
    )
    add_para(doc, "（3）升力系数与阻力系数（α 为攻角，弧度制参与三角函数运算）：")
    add_para(doc, "Cl = Cn·cosα − Ca·sinα", indent=False)
    add_para(doc, "Cd = Cn·sinα + Ca·cosα", indent=False)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["攻角 α/(°)", "Cn", "Ca", "Cl", "Cd", "Cl/Cd"]):
        hdr[i].text = h
    for _, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = f"{row['alpha_deg']:.1f}"
        cells[1].text = f"{row['Cn']:.4f}"
        cells[2].text = f"{row['Ca']:.4f}"
        cells[3].text = f"{row['Cl']:.4f}"
        cells[4].text = f"{row['Cd']:.4f}"
        ld = row["L_D"]
        cells[5].text = f"{ld:.2f}" if abs(ld) < 50 else "—"

    cap = doc.add_paragraph("表1  各攻角下力系数计算结果（V = 20 m/s）")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "5.4 升阻特性曲线及分析", level=2)
    add_image(doc, "图3_升阻特性曲线.png", "图3  升力系数与阻力系数随攻角变化曲线")
    add_image(doc, "图4_升力阻力极曲线.png", "图4  升力-阻力极曲线", width_cm=10)
    add_image(doc, "图5_升阻比曲线.png", "图5  升阻比随攻角变化曲线")

    add_para(
        doc,
        "（1）升力特性：由表1和图3，在 α = −2°～12° 范围内，"
        "升力系数 Cl 随攻角单调增大，由 −0.14 增至 0.97，"
        "近似满足对称翼型小攻角线性关系 Cl ≈ 2π·α（α 为弧度）。"
        "例如 α = 8°（0.140 rad）时，理论值 2πα ≈ 0.88，实验值 0.60，"
        "偏低原因包括：NACA0025 厚度达 25%，偏离薄翼假设；"
        "测压孔有限、积分误差；风洞来流不完全均匀等。"
    )
    add_para(
        doc,
        "（2）阻力特性：零攻角时 Cd ≈ 0.053，来自厚翼型型阻及压差阻力。"
        "随 α 增大，Cd 持续上升（0° 时 0.053 → 16° 时 0.390），"
        "主要因为迎风投影增大、尾缘压差及流动分离区扩大，压差阻力显著增加。"
    )
    add_para(
        doc,
        "（3）失速与升阻比：α = 16° 时 Cl 由 12° 的 0.97 降至 0.93，"
        "而 Cd 由 0.32 增至 0.39，说明翼型已进入失速区——"
        "上表面大面积分离使有效环量下降，升力不再增加反而略降，阻力急剧增大。"
        "升阻比 Cl/Cd 在 α ≈ 8° 时达到最大（约 3.7），"
        "是兼顾升力与阻力的较优工作点；大攻角下升阻比迅速恶化。"
    )
    add_para(
        doc,
        "（4）极曲线（图4）：Cl–Cd 曲线呈典型“极曲线”形态，"
        "左下方对应小攻角、小升力，沿曲线逆时针方向对应 α 增大；"
        "曲线左上端对应接近失速的最大升力工况。"
    )

    add_heading(doc, "5.5 延伸思考题", level=2)

    add_para(doc, "（1）升力系数是否随攻角的增大一直保持增大？为什么？", indent=False)
    add_para(
        doc,
        "答：不是。升力系数随攻角的变化可分为三个阶段："
    )
    add_para(
        doc,
        "① 线性段（未失速区）：在小攻角范围内，翼型上下表面压差随 α 近似线性增大，"
        "环量 Γ 与 sinα 或 α（小角度）成正比，故 Cl 近似线性增大。"
        "本实验中 −2°～12° 即处于该阶段，Cl 从 −0.14 增至 0.97。"
        "薄翼理论给出 Cl = 2πα，对称翼型零升攻角为 0°，与本实验趋势一致。"
    )
    add_para(
        doc,
        "② 接近临界攻角：当 α 继续增大，上表面边界层沿翼面发展，"
        "在逆压梯度区可能出现局部分离，Cl 增速放缓。"
        "本实验 12° 时 Cl 已达 0.97，接近 NACA0025 在该 Re 下的最大升力系数。"
    )
    add_para(
        doc,
        "③ 失速段：超过临界攻角 αcr 后，上表面气流大面积分离，"
        "前缘吸力峰崩溃，有效迎流面积和环量下降，Cl 不再增大甚至减小，"
        "同时分离区增大涡阻，Cd 急剧上升。本实验 16° 时 Cl 降至 0.93、"
        "Cd 升至 0.39，即出现典型失速特征。若继续增大攻角，"
        "Cl 将大幅下降，翼型进入深度失速。"
    )
    add_para(
        doc,
        "物理本质：升力来源于翼型对流体的环量作用；"
        "失速的本质是边界层无法克服逆压梯度而发生分离，"
        "破坏了翼型设计所期望的压力分布，因此升力不再随 α 单调增加。"
    )

    add_para(doc, "（2）流体速度对升阻特性的影响？", indent=False)
    add_para(
        doc,
        "答：需从无量纲系数与有量纲气动力、以及雷诺数效应两方面分析："
    )
    add_para(
        doc,
        "① 对 Cl、Cd 无量纲系数的影响（低速不可压）："
        "在 Ma ≪ 1 时，由量纲分析，Cl 和 Cd 主要决定于攻角 α 和翼型几何，"
        "理论上与来流速度 V 无直接函数关系。"
        "即同一翼型在同一 α 下，V 从 10 m/s 变到 30 m/s，"
        "Cl、Cd 应基本不变（实验测量的微小差异来自 Re 变化，见下）。"
        "有量纲升力 L = Cl·q∞·S = Cl·½ρV²·S，故 V 增大时升力按 V² 增大。"
    )
    add_para(
        doc,
        f"② 雷诺数效应：Re = ρVL/μ 与速度成正比。"
        f"本实验 V = 20 m/s、c = 0.2 m、ρ = {RHO:.3f} kg/m³ 时 Re ≈ {RE:.2e}。"
        f"Re 影响边界层流态：Re 较小时层流边界层更易分离，αcr 降低，Clmax 减小；"
        f"Re 较大时流动更稳定，分离点后移，αcr 和 Clmax 可略有提高。"
        f"若将 V 减半至 10 m/s，Re 减半，可能更早失速；"
        f"若 V 增至 40 m/s，Re 加倍，失速攻角可能略推迟，但仍在同一量级。"
    )
    add_para(
        doc,
        "③ 摩擦阻力：总阻力 Cd = 压差阻力 + 摩擦阻力。"
        "摩擦阻力与 Re 有关：Re 增大时边界层相对更薄，摩擦阻力系数略降；"
        "但本实验以压差阻力为主（厚翼型、较大 α），速度对 Cd 的影响不如对 Re 间接影响显著。"
    )
    add_para(
        doc,
        "④ 压缩性效应：当 V 增大使 Ma > 0.3 时，空气密度变化、"
        "可能出现局部超音速区和激波，产生波阻，Cl–α 曲线斜率改变，"
        "升阻特性与低速规律显著不同。本实验 V = 20 m/s，"
        "远低于声速，可忽略压缩性。"
    )
    add_para(
        doc,
        "⑤ 实验建议：若要研究速度影响，应在多个 V 下重复同一组 α，"
        "绘制不同 Re 下的 Cl–α 曲线对比；同时记录各次实验的 P、t，"
        "准确计算 ρ 和 Re，以保证数据可比性。"
    )

    out = os.path.join(DESKTOP, "实验报告_第4-5部分_翼型风洞.docx")
    doc.save(out)
    print(f"Saved: {out}")
    print(f"P={P_ATM} Pa, t={T_C} C, rho={RHO:.3f} kg/m3, q={Q_INF:.1f} Pa, Re={RE:.2e}")


if __name__ == "__main__":
    main()
